from dataclasses import asdict
from datetime import datetime
import io
import re
import zipfile

from flask import redirect, url_for, Response, request, jsonify, send_file
from werkzeug.exceptions import HTTPException
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

import app as app_module

app = app_module.app
status_payload = app_module.status_payload
parse_lessons_from_request = app_module.parse_lessons_from_request
logger = app_module.logger
check_usage_limit = app_module.check_usage_limit
store_docx_file = app_module.store_docx_file
clean_text = app_module.clean_text

RLM = "\u200f"
LRM = "\u200e"


@app.route('/favicon.ico')
def favicon():
    return Response(status=204)


@app.route('/generate', methods=['GET'], endpoint='generate_get')
def generate_get():
    return redirect(url_for('index'))


def _ascii_name(prefix='Lesson_Plan', ext='docx'):
    return f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"


def subject_family(subject: str) -> str:
    s = (subject or '').lower().strip()
    if any(x in s for x in ['لغة عربية', 'اللغه العربيه', 'عربي', 'arabic']):
        return 'arabic'
    if any(x in s for x in ['رياض', 'math', 'calculus', 'جبر', 'هندسة']):
        return 'math'
    if any(x in s for x in ['science', 'علوم', 'physics', 'chemistry', 'biology', 'فيزياء', 'كيمياء', 'أحياء']):
        return 'science'
    if any(x in s for x in ['english', 'انجليزي', 'لغة انجليزية']):
        return 'english'
    return 'general'


def topic_key(topic: str, subject: str) -> str:
    t = (topic or '').lower()
    sf = subject_family(subject)
    if sf == 'arabic':
        if any(x in t for x in ['النعت', 'نعت', 'الصفة', 'صفة']):
            return 'arabic_naat'
        if any(x in t for x in ['المضاف', 'الإضافة', 'مضاف إليه']):
            return 'arabic_idafa'
        if any(x in t for x in ['كان وأخواتها', 'إن وأخواتها', 'كان', 'إن']):
            return 'arabic_grammar'
        if any(x in t for x in ['قراءة', 'نص', 'قصة', 'قصيدة', 'شعر']):
            return 'arabic_reading'
        return 'arabic_general'
    if any(x in t for x in ['طول المنحنى', 'arc length', 'curve length']) and any(x in t for x in ['مماس', 'مماسات', 'tangent']):
        return 'math_tangent_arc'
    if any(x in t for x in ['مماس', 'مماسات', 'tangent']):
        return 'math_tangent'
    if any(x in t for x in ['طول المنحنى', 'arc length']):
        return 'math_arc'
    if any(x in t for x in ['اشتقاق', 'مشتقة', 'derivative']):
        return 'math_derivative'
    if any(x in t for x in ['نهاية', 'نهايات', 'limit', 'اتصال']):
        return 'math_limits'
    return f'{sf}_general'


def bullets(items, lang='ar'):
    if lang == 'ar':
        return '\n'.join(f'{RLM}{i}. {item}' for i, item in enumerate(items, 1))
    return '\n'.join(f'{i}. {item}' for i, item in enumerate(items, 1))


def m(expr: str) -> str:
    """Readable Word-friendly math, not raw LaTeX."""
    e = expr.strip()
    replacements = {
        "f'(a)": 'f′(a)', "f'(x)": 'f′(x)', "f'(1)": 'f′(1)', "f'(2)": 'f′(2)',
        'x^2': 'x²', '(f′(x))^2': '(f′(x))²', '2x': '2x',
        'y-f(a)=f′(a)(x-a)': 'y − f(a) = f′(a)(x − a)',
        'y-5=4(x-2)': 'y − 5 = 4(x − 2)',
        'f(x)=x²+1': 'f(x) = x² + 1',
        'f(x)=x²-3x': 'f(x) = x² − 3x',
        'm=f′(a)': 'm = f′(a)',
        'm=f′(2)=4': 'm = f′(2) = 4',
        'f(2)=5': 'f(2) = 5',
        'x=2': 'x = 2', 'x=1': 'x = 1',
        '[a,b]': '[a, b]',
    }
    e = e.replace('\\int_a^b\\sqrt{1+(f\'(x))^2}\\,dx', 'L = ∫ₐᵇ √(1 + (f′(x))²) dx')
    e = e.replace('\\int_a^b\\sqrt{1+(f′(x))^2}\\,dx', 'L = ∫ₐᵇ √(1 + (f′(x))²) dx')
    e = e.replace('L=', 'L = ')
    e = e.replace('^2', '²').replace('^3', '³').replace('^4', '⁴')
    e = e.replace("f'", 'f′')
    for k, v in replacements.items():
        e = e.replace(k, v)
    e = re.sub(r'\s+', ' ', e).strip()
    return f'{LRM}{e}{LRM}'


def topic_examples(topic: str, subject: str, lang: str):
    key = topic_key(topic, subject)
    if lang == 'ar':
        if key == 'arabic_naat':
            return {
                'keywords': 'النعت، المنعوت، المطابقة، تابع، الإعراب، التعريف والتنكير، النوع والعدد',
                'concept': 'النعت تابع يصف اسمًا قبله يسمى المنعوت، ويطابقه في الإعراب والتعريف أو التنكير والنوع والعدد. يفرق الطلاب بين النعت والخبر من خلال وظيفة الكلمة وموقعها في الجملة.',
                'worked': 'مثال محلول: «جاءَ الطالبُ المجتهدُ». كلمة «المجتهدُ» نعت مرفوع؛ لأنها وصفت «الطالبُ» ووافقته في الرفع والتعريف والتذكير والإفراد. مثال آخر: «قرأتُ قصةً ممتعةً»؛ «ممتعةً» نعت منصوب يصف «قصةً» ويوافقه في النصب والتنكير والتأنيث والإفراد.',
                'guided': 'تدريب موجه: استخرج النعت والمنعوت وبيّن وجه المطابقة في: «كرّمت المدرسةُ الطالباتِ المتميزاتِ»، «شاهدتُ منظرًا جميلًا»، «مررتُ بمعلمٍ مبدعٍ». ثم غيّر المنعوت من مفرد إلى مثنى وجمع وعدّل النعت بدقة.',
                'hots': 'سؤال تفكير عليا: قارن بين «الطالبُ مجتهدٌ» و«الطالبُ المجتهدُ حاضرٌ». لماذا كانت «مجتهدٌ» خبرًا في الأولى بينما «المجتهدُ» نعتًا في الثانية؟',
                'misconception': 'الخلط بين النعت والخبر، أو عدم مطابقة النعت للمنعوت في التعريف والتنكير أو علامة الإعراب.'
            }
        if key == 'math_tangent_arc':
            return {
                'keywords': 'المماس، ميل المماس، المشتقة، طول المنحنى، معدل التغير اللحظي، التكامل المحدد',
                'concept': f'يربط الدرس بين المشتقة كمعدل تغير لحظي وطول المنحنى كقيمة تراكمية للمسافة. الصيغ الرئيسة: {m("m=f'(a)")}، {m("y-f(a)=f'(a)(x-a)")}، {m("L=\\int_a^b\\sqrt{1+(f'(x))^2}\\,dx")}.',
                'worked': f'مثال محلول: إذا كانت {m("f(x)=x^2+1")} عند {m("x=2")} فإن {m("f(2)=5")} و {m("f'(x)=2x")}، لذلك {m("m=f'(2)=4")} ومعادلة المماس {m("y-5=4(x-2)")}. ثم نوضح أن طول المنحنى لا يساوي المسافة المستقيمة بل يعتمد على {m("L=\\int_a^b\\sqrt{1+(f'(x))^2}\\,dx")}.',
                'guided': f'تدريب موجه: للدالة {m("f(x)=x^2-3x")} عند {m("x=1")} أوجد {m("f'(1)")} واكتب معادلة المماس، ثم فسّر لماذا تظهر المشتقة داخل صيغة طول المنحنى.',
                'hots': f'سؤال إثرائي: قارن بين ميل المماس عند نقطة وطول المنحنى على الفترة {m("[a,b]")}. أيهما لحظي وأيهما تراكمي؟',
                'misconception': 'الخلط بين قيمة الدالة وقيمة المشتقة، أو اعتبار طول المنحنى مساويًا للمسافة المستقيمة بين الطرفين.'
            }
        # generic Arabic by subject
        return {
            'keywords': f'{topic}، مفاهيم أساسية، تطبيق، تحليل، تقويم',
            'concept': f'يركز درس {topic} في مادة {subject} على فهم المفهوم من خلال مثال واضح، ثم تطبيق موجه، ثم مهمة مستقلة تكشف الفهم الحقيقي لا الحفظ.',
            'worked': f'مثال محلول خاص بدرس {topic}: يعرض المعلم موقفًا قصيرًا أو مسألة مباشرة، يحدد الكلمات المفتاحية، ثم يشرح خطوات التفكير ويبرز سبب اختيار الإجابة.',
            'guided': f'تدريب موجه: يطبق الطلاب فكرة {topic} في مثال جديد، ثم يكتبون جملة تفسيرية توضح لماذا كانت الإجابة صحيحة.',
            'hots': f'سؤال تفكير عليا: صمّم مثالًا جديدًا على {topic} أو غيّر شرطًا واحدًا، ثم ناقش أثر التغيير في الإجابة.',
            'misconception': f'خطأ شائع في درس {topic}: الاكتفاء بالحفظ دون ربط القاعدة بالمثال أو الدليل.'
        }
    return {
        'keywords': f'{topic}, concept, application, analysis, assessment',
        'concept': f'The lesson builds understanding of {topic} in {subject} through a model, guided task, and independent evidence of learning.',
        'worked': f'Worked example linked directly to {topic} with clear reasoning and key vocabulary.',
        'guided': f'Guided practice on {topic} followed by an interpretation sentence.',
        'hots': f'HOTS: create a new example or change one condition and explain the impact.',
        'misconception': f'Common misconception in {topic}: memorising the rule without using evidence or reasoning.'
    }


def stable_content(lesson):
    lang = lesson.language
    topic = (lesson.topic or ('الدرس' if lang == 'ar' else 'the lesson')).strip()
    subject = (lesson.subject or ('رياضيات' if lang == 'ar' else 'Mathematics')).strip()
    class_name = (lesson.class_name or ('الثاني عشر متقدم' if lang == 'ar' else 'Grade 12 Advanced')).strip()
    ex = topic_examples(topic, subject, lang)
    sf = subject_family(subject)
    note = (lesson.notes or '').strip()

    if lang == 'ar':
        is_arabic = sf == 'arabic'
        sdg = 'SDG 4 التعليم الجيد: تنمية الكفاءة اللغوية والتواصل الفعال والاعتزاز باللغة العربية.' if is_arabic else 'SDG 4 التعليم الجيد + SDG 11: توظيف المعرفة في اتخاذ قرارات دقيقة ومسؤولة.'
        identity = 'الهوية الوطنية: ربط الدرس باللغة العربية بوصفها وعاء الهوية والثقافة في دولة الإمارات.' if is_arabic else 'الهوية الوطنية والاستدامة: ربط الدقة والانضباط بثقافة التميز في دولة الإمارات.'
        resources = 'نصوص قصيرة، بطاقات كلمات، سبورة ذكية، أمثلة تحليلية، دفتر الطالب، بطاقة خروج.' if is_arabic else 'سبورة ذكية، ورقة عمل قصيرة، بطاقات خطوات، رسم/جدول، Classroom Monitor، دفتر الطالب.'
        return {
            'subject': subject,
            'class_name': class_name,
            'keywords': ex['keywords'],
            'sdg': sdg,
            'strategies': 'تمهيد تشخيصي قصير، نموذج محلول، تفكير بصوت عالٍ، تدريب موجه، تطبيق مستقل، سؤال تفكير عليا، وتغذية راجعة فورية.' + (f'\nملاحظة المعلم: {note}' if note else ''),
            'intervention': 'دعم فوري: بطاقة خطوات، مثال جزئي، سؤال تحقق قصير بعد كل خطوة، وشريك داعم. إذا أخفق أكثر من 25% في مهمة AFL تُنفذ إعادة تدريس مركزة.\nخطأ متوقع: ' + ex['misconception'],
            'learning_outcomes': bullets([
                f'أفسر مفهوم {topic} باستخدام مصطلحات دقيقة من مادة {subject}.',
                'أميز الفكرة المستهدفة من أمثلة صحيحة وأخرى خاطئة.',
                'أطبق القاعدة أو المهارة في مثال مباشر مع تبرير الخطوات.',
                'أحل نشاطًا متدرجًا مرتبطًا بعنوان الدرس.',
                'أقارن بين حالتين وأفسر الفرق بدليل واضح.',
                'أصحح خطأً شائعًا وأكتب قاعدة مختصرة للتمييز.'
            ], 'ar'),
            'differentiation': bullets([
                'دعم: بطاقة خطوات، مثال جزئي، كلمات مفتاحية ملوّنة، وتقليل الحمل الكتابي عند الحاجة.',
                'مستوى متوقع: تدريب موجه ثم تطبيق مستقل مشابه للنموذج.',
                'متقدمون: سؤال HOTS يتطلب مقارنة أو إنتاج مثال جديد مع تبرير.',
                'IEP/APL: تبسيط الصياغة، وقت إضافي، وتمثيل بصري عند الحاجة.'
            ], 'ar'),
            'success_criteria': bullets([
                f'أشرح فكرة {topic} بجملة دقيقة.',
                'أحدد العناصر أو الكلمات المفتاحية دون خلط.',
                'أطبق المهارة في مثال جديد.',
                'أبرر إجابتي بدليل واضح.',
                'أصحح خطأً شائعًا وأوضح السبب.',
                'أحقق 80% فأكثر في بطاقة الخروج أو أكتب خطوة تحسين.'
            ], 'ar'),
            'starter': f'نشاط تمهيدي (5-7 دقائق): يعرض المعلم مثالين مرتبطين بدرس {topic}؛ أحدهما صحيح والآخر يتضمن خطأ شائعًا. يحدد الطلاب الفرق ويبررون إجابتهم.\n{ex["guided"]}',
            'main': 'أنشطة رئيسية:\n' + bullets([ex['worked'], ex['guided'], 'تطبيق فردي قصير ثم مقارنة زوجية للحل أو التحليل.', ex['hots']], 'ar'),
            'teacher_led': f'دور المعلم: يشرح المفهوم من مثال واضح، يبرز الكلمات المفتاحية/الصيغ، ويسأل أسئلة تحقق قصيرة.\n{ex["concept"]}',
            'student_led': 'دور الطلاب: يحلون تدريبًا موجهًا ثم مهمة مستقلة، ويشرح كل طالب خطوة أو سببًا لزميله، ثم يكتب جملة تفسيرية قصيرة.',
            'plenary': 'خاتمة وتقويم: بطاقة خروج من 3 أجزاء: تحديد المفهوم، تطبيق قصير، وتصحيح خطأ شائع. يعرض المعلم إجابة نموذجية وخطوة تحسين.',
            'kpi': 'KPI AFL Task: 4 أسئلة قصيرة: مفهوم، تطبيق مباشر، تفسير، وتصحيح خطأ. معيار النجاح 80% فأكثر مع تدخل فوري.',
            'resources': resources,
            'identity': identity,
            'competency': 'تواصل، تفكير ناقد، حل مشكلات، تعاون، إبداع، مسؤولية ذاتية، ووعي رقمي.',
            'curriculum': f'ارتباط المنهج: {subject} - {class_name} - {topic}. يتكامل مع الفهم، التطبيق، التفسير، التقويم، والتواصل الشفهي والكتابي.',
            '_mode': 'professional_subject_specific_compact'
        }

    return {
        'subject': subject, 'class_name': class_name, 'keywords': ex['keywords'],
        'sdg': 'SDG 4 Quality Education: applying subject knowledge through reasoning and communication.',
        'strategies': 'Diagnostic starter, worked example, think-aloud modelling, guided practice, independent task, HOTS question, and immediate feedback.',
        'intervention': 'Support: step card, partial example, check questions, and peer support. Misconception: ' + ex['misconception'],
        'learning_outcomes': bullets([f'Explain {topic} using accurate {subject} terminology.', 'Identify the target idea from correct and incorrect examples.', 'Apply the skill in a direct example with justification.', 'Complete a progressive task linked to the title.', 'Compare two cases using evidence.', 'Correct a common misconception.'], 'en'),
        'differentiation': bullets(['Support: step card and partial example.', 'Expected: guided then independent task.', 'Advanced: HOTS comparison or new example.', 'IEP/APL: simplified wording and extra time.'], 'en'),
        'success_criteria': bullets([f'I can explain {topic}.', 'I can identify key elements.', 'I can apply the skill.', 'I can justify my answer.', 'I can correct a common error.', 'I can reach 80% in the exit ticket.'], 'en'),
        'starter': f'Starter: compare two examples linked to {topic}.\n{ex["guided"]}',
        'main': 'Main activities:\n' + bullets([ex['worked'], ex['guided'], 'Independent application and peer comparison.', ex['hots']], 'en'),
        'teacher_led': f'Teacher models and checks understanding.\n{ex["concept"]}',
        'student_led': 'Students complete guided practice, solve independently, and explain reasoning to a peer.',
        'plenary': 'Exit Ticket: concept, application, and error correction.',
        'kpi': 'KPI AFL Task: concept, application, interpretation, and error correction. Benchmark: 80%.',
        'resources': 'Smart board, short worksheet, step cards, Classroom Monitor, student notebook.',
        'identity': 'Connect learning to excellence, respect, and responsible participation.',
        'competency': 'Communication, critical thinking, problem solving, collaboration, creativity, self-management.',
        'curriculum': f'Curriculum link: {subject} - {class_name} - {topic}.',
        '_mode': 'professional_subject_specific_en'
    }


app_module.build_content = stable_content


def improved_set_cell_text(cell, text: str, lang: str = 'en', size: float = 8.0, bold: bool = False) -> None:
    rtl = lang == 'ar'
    font = 'Arial' if rtl else 'Times New Roman'
    cell.text = ''
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    lines = clean_text(text).split('\n') if text else ['']
    for i, line in enumerate(lines):
        line = line.strip()
        if rtl and re.match(r'^\d+[\.)]\s+', line):
            line = RLM + line
        paragraph = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        app_module.set_paragraph_bidi(paragraph, rtl)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        app_module.set_run_font(run, font, max(size + 0.45, 7.3 if rtl else 7.0), bold=True)


app_module.set_cell_text = improved_set_cell_text


def normalize_docx(docx_bytes: bytes) -> bytes:
    """Remove fixed row heights to stop huge blanks, and keep section titles with their following content."""
    doc = Document(io.BytesIO(docx_bytes))
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            for h in list(tr_pr.findall(qn('w:trHeight'))):
                tr_pr.remove(h)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    if 'Lesson Structure' in p.text or 'LESSON PLAN' in p.text:
                        p.paragraph_format.keep_with_next = True
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


_original_generate_docx = app_module.generate_docx


def generate_docx_clean(lesson):
    data = _original_generate_docx(lesson)
    return normalize_docx(data)


app_module.generate_docx = generate_docx_clean


def fallback_docx(lesson, reason=''):
    content = stable_content(lesson)
    doc = Document()
    doc.add_heading('Magdy Lesson Planner - Lesson Plan', 0)
    for title_text, key in [('Learning Outcomes', 'learning_outcomes'), ('Success Criteria', 'success_criteria'), ('Starter', 'starter'), ('Main Activities', 'main'), ('Teacher-led', 'teacher_led'), ('Student-led', 'student_led'), ('Plenary', 'plenary')]:
        p = doc.add_paragraph(); r = p.add_run(title_text); r.bold = True
        doc.add_paragraph(content.get(key, ''))
    out = io.BytesIO(); doc.save(out); return out.getvalue()


def safe_preview():
    try:
        lessons, errors = parse_lessons_from_request()
        if errors:
            return jsonify({'ok': False, 'errors': errors, 'status': status_payload()}), 400
        lesson = lessons[0]
        return jsonify({'ok': True, 'lesson': asdict(lesson), 'content': stable_content(lesson), 'status': status_payload()})
    except Exception as exc:
        logger.exception('Safe preview failed')
        return jsonify({'ok': False, 'error': 'تعذر إنشاء المعاينة. تم تسجيل الخطأ في سجل التطبيق.', 'details': str(exc), 'status': status_payload()}), 500


def safe_generate():
    try:
        lessons, errors = parse_lessons_from_request()
        if errors:
            return Response(' | '.join(errors), status=400, mimetype='text/plain; charset=utf-8')
        user_key = lessons[0].teacher or request.remote_addr or 'anonymous'
        files = []
        for lesson in lessons:
            ok_limit, limit_msg = check_usage_limit(user_key)
            if not ok_limit:
                return Response(limit_msg, status=429, mimetype='text/plain; charset=utf-8')
            try:
                docx_bytes = app_module.generate_docx(lesson)
                try:
                    store_docx_file(lesson, docx_bytes)
                except Exception:
                    logger.exception('Could not store generated file in library; continuing download')
            except Exception as exc:
                logger.exception('Official template generation failed; using fallback DOCX')
                docx_bytes = fallback_docx(lesson, str(exc))
            files.append((_ascii_name(f'Lesson_Plan_{lesson.index}', 'docx'), docx_bytes))
        if len(files) == 1:
            filename, docx_bytes = files[0]
            return send_file(io.BytesIO(docx_bytes), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename, max_age=0)
        zip_bytes = io.BytesIO()
        with zipfile.ZipFile(zip_bytes, 'w', zipfile.ZIP_DEFLATED) as zf:
            for filename, docx_bytes in files:
                zf.writestr(filename, docx_bytes)
        zip_bytes.seek(0)
        return send_file(zip_bytes, mimetype='application/zip', as_attachment=True, download_name=_ascii_name('Lesson_Plans_Batch', 'zip'), max_age=0)
    except Exception as exc:
        logger.exception('Safe generate failed')
        return Response(f'Internal generation error. Application log details: {exc}', status=500, mimetype='text/plain; charset=utf-8')


app.view_functions['preview'] = safe_preview
app.view_functions['generate'] = safe_generate


@app.errorhandler(HTTPException)
def handle_http_exception(exc):
    if request.path.startswith('/api/'):
        return jsonify({'ok': False, 'error': exc.description, 'status_code': exc.code, 'status': status_payload()}), exc.code
    if exc.code in (404, 405):
        return redirect(url_for('index'))
    return exc
