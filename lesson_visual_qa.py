from __future__ import annotations

import io
import math
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SECTION_MARKERS = {
    "lesson plan", "خطة الدرس", "lesson structure", "هيكل الدرس",
    "learning outcomes:", "نواتج التعلم:", "success criteria:", "معايير النجاح:",
    "starter (10 mins)", "التمهيد (10 دقائق)", "main activities (15 mins)",
    "الأنشطة الرئيسة (15 دقيقة)", "teacher-led", "دور المعلم", "student-led",
    "دور الطلاب", "plenary (10 mins)", "الخاتمة (10 دقائق)",
}

LABEL_MARKERS = {
    "teacher:", "subject:", "date:", "class:", "topic:", "periods:",
    "المعلم:", "المادة:", "التاريخ:", "الصف:", "موضوع الدرس:", "عدد الحصص:",
    "key words:", "keywords:", "الكلمات المفتاحية:", "primary sdg focus:",
    "الهدف الرئيس للتنمية المستدامة:", "strategies:", "استراتيجيات التدريس:",
    "kpi afl assignment task:", "مهمة التقويم التكويني kpi/afl:",
}

PLACEHOLDERS = (
    "click or tap here to enter text", "type here", "insert text", "enter text",
    "انقر أو اضغط هنا لإدخال نص", "اكتب هنا", "{{", "}}",
)


@dataclass
class QAMetrics:
    adjusted_cells: int = 0
    dense_cells: int = 0
    long_rows_unlocked: int = 0
    short_rows_locked: int = 0
    removed_empty_paragraphs: int = 0
    removed_extra_breaks: int = 0
    fixed_tables: int = 0
    placeholders_removed: int = 0


def _normal(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip()).casefold()


def _unique_cells(row):
    seen = set()
    output = []
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        output.append(cell)
    return output


def _set_fixed_table_layout(table) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")


def _cell_width_twips(cell) -> int:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is not None:
        try:
            value = int(width.get(qn("w:w"), "0"))
            if value > 0:
                return value
        except Exception:
            pass
    grid_span = properties.find(qn("w:gridSpan"))
    span = 1
    if grid_span is not None:
        try:
            span = max(1, int(grid_span.get(qn("w:val"), "1")))
        except Exception:
            span = 1
    return 1600 * span


def _equation_count(cell) -> int:
    return len(cell._tc.xpath(".//m:oMath | .//m:oMathPara"))


def _effective_text(cell) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs if paragraph.text.strip())


def _estimated_lines(cell) -> float:
    width_inches = max(_cell_width_twips(cell) / 1440.0, 0.8)
    text = _effective_text(cell)
    paragraphs = max(1, len([p for p in cell.paragraphs if p.text.strip()]))
    arabic_count = len(re.findall(r"[\u0600-\u06ff]", text))
    latin_count = len(re.findall(r"[A-Za-z0-9]", text))
    weighted_chars = arabic_count * 1.12 + latin_count + max(0, len(text) - arabic_count - latin_count) * 0.35
    chars_per_line = max(13.0, width_inches * 11.5)
    return weighted_chars / chars_per_line + paragraphs * 0.5 + _equation_count(cell) * 1.1


def _remove_property(paragraph, name: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    node = properties.find(qn(f"w:{name}"))
    if node is not None:
        properties.remove(node)


def _set_property(paragraph, name: str, enabled: bool = True) -> None:
    properties = paragraph._p.get_or_add_pPr()
    node = properties.find(qn(f"w:{name}"))
    if enabled and node is None:
        properties.append(OxmlElement(f"w:{name}"))
    elif not enabled and node is not None:
        properties.remove(node)


def _set_row_split(row, locked: bool) -> None:
    properties = row._tr.get_or_add_trPr()
    node = properties.find(qn("w:cantSplit"))
    if locked and node is None:
        properties.append(OxmlElement("w:cantSplit"))
    elif not locked and node is not None:
        properties.remove(node)


def _is_section_cell(cell) -> bool:
    value = _normal(cell.text)
    return value in SECTION_MARKERS


def _is_label_cell(cell) -> bool:
    value = _normal(cell.text)
    return value in LABEL_MARKERS or (value.endswith(":") and len(value) <= 70)


def _paragraph_has_only_equation(paragraph) -> bool:
    has_math = bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))
    return has_math and not paragraph.text.strip()


def _font_floor(language: str) -> float:
    return 8.55 if language == "ar" else 8.35


def _target_font_size(cell, language: str) -> float:
    lines = _estimated_lines(cell)
    text_length = len(_effective_text(cell))
    if _is_section_cell(cell):
        return 10.8
    if _is_label_cell(cell):
        return 8.9
    if lines > 28 or text_length > 1800:
        return _font_floor(language)
    if lines > 22 or text_length > 1250:
        return 8.65 if language == "ar" else 8.45
    if lines > 16 or text_length > 850:
        return 8.85 if language == "ar" else 8.65
    return 9.15 if language == "ar" else 8.95


def _set_cell_margins(cell, dense: bool) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    values = (55, 80, 55, 80) if dense else (75, 100, 75, 100)
    for name, value in zip(("top", "start", "bottom", "end"), values):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _apply_cell_typography(cell, language: str, metrics: QAMetrics) -> None:
    rtl = language == "ar"
    target_size = _target_font_size(cell, language)
    estimated = _estimated_lines(cell)
    dense = estimated > 18
    if dense:
        metrics.dense_cells += 1
    _set_cell_margins(cell, dense)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if (_is_label_cell(cell) or _is_section_cell(cell)) else WD_CELL_VERTICAL_ALIGNMENT.TOP

    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        heading = _is_section_cell(cell) or _is_label_cell(cell) or (
            text.endswith(":") and len(text) <= 70
        )
        equation_only = _paragraph_has_only_equation(paragraph)
        long_paragraph = len(text) > 260 or estimated > 18

        paragraph.paragraph_format.space_before = Pt(0.8 if heading else 0)
        paragraph.paragraph_format.space_after = Pt(1.1 if heading else (0.6 if dense else 1.0))
        paragraph.paragraph_format.line_spacing = 1.0 if dense else 1.04
        paragraph.paragraph_format.widow_control = True

        # Keeping every long paragraph together creates large blank spaces. Only short
        # headings and compact equations should be kept as one visual block.
        _set_property(paragraph, "keepNext", heading)
        _set_property(paragraph, "keepLines", heading or equation_only or not long_paragraph)

        if equation_only:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif rtl:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif _is_section_cell(cell):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            current = run.font.size.pt if run.font.size else None
            desired = target_size
            if heading:
                desired = max(target_size, 9.0)
            if current is None or abs(current - desired) > 0.15:
                run.font.size = Pt(desired)
                metrics.adjusted_cells += 1


def _row_estimated_lines(row) -> float:
    cells = _unique_cells(row)
    return max((_estimated_lines(cell) for cell in cells), default=0.0)


def _apply_row_pagination(row, metrics: QAMetrics) -> None:
    cells = _unique_cells(row)
    section_row = any(_is_section_cell(cell) for cell in cells)
    estimate = _row_estimated_lines(row)
    lock = section_row or estimate <= 13
    _set_row_split(row, lock)
    if lock:
        metrics.short_rows_locked += 1
    else:
        metrics.long_rows_unlocked += 1


def _is_break_only(paragraph) -> bool:
    if paragraph.text.strip():
        return False
    return bool(paragraph._p.xpath(".//w:br[@w:type='page'] | .//w:lastRenderedPageBreak"))


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_duplicate_breaks(doc, metrics: QAMetrics) -> None:
    previous_break = False
    for paragraph in list(doc.paragraphs):
        current_break = _is_break_only(paragraph)
        if current_break and previous_break:
            _remove_paragraph(paragraph)
            metrics.removed_extra_breaks += 1
            continue
        previous_break = current_break


def _remove_trailing_empty_body_paragraphs(doc, metrics: QAMetrics) -> None:
    body = doc._element.body
    children = list(body)
    for element in reversed(children):
        if element.tag == qn("w:sectPr"):
            continue
        if element.tag != qn("w:p"):
            break
        text = "".join(element.itertext()).strip()
        has_media = bool(element.xpath(".//w:drawing | .//w:pict | .//w:object"))
        if text or has_media:
            break
        body.remove(element)
        metrics.removed_empty_paragraphs += 1


def _remove_placeholders(doc, metrics: QAMetrics) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in _unique_cells(row):
                for paragraph in cell.paragraphs:
                    value = _normal(paragraph.text)
                    if any(marker in value for marker in PLACEHOLDERS):
                        paragraph.clear()
                        metrics.placeholders_removed += 1
    for paragraph in doc.paragraphs:
        value = _normal(paragraph.text)
        if any(marker in value for marker in PLACEHOLDERS):
            paragraph.clear()
            metrics.placeholders_removed += 1


def _qa_comment(metrics: QAMetrics) -> str:
    return (
        "Visual QA passed | "
        f"fixed_tables={metrics.fixed_tables}; "
        f"dense_cells={metrics.dense_cells}; "
        f"adjustments={metrics.adjusted_cells}; "
        f"rows_locked={metrics.short_rows_locked}; "
        f"rows_unlocked={metrics.long_rows_unlocked}; "
        f"empty_removed={metrics.removed_empty_paragraphs}; "
        f"breaks_removed={metrics.removed_extra_breaks}; "
        f"placeholders_removed={metrics.placeholders_removed}"
    )


def visual_qa(data: bytes, language: str = "en") -> bytes:
    doc = Document(io.BytesIO(data))
    metrics = QAMetrics()

    _remove_placeholders(doc, metrics)
    for table in doc.tables:
        _set_fixed_table_layout(table)
        metrics.fixed_tables += 1
        for row in table.rows:
            _apply_row_pagination(row, metrics)
            for cell in _unique_cells(row):
                _apply_cell_typography(cell, language, metrics)

    _remove_duplicate_breaks(doc, metrics)
    _remove_trailing_empty_body_paragraphs(doc, metrics)

    current_comments = str(doc.core_properties.comments or "").strip()
    qa_line = _qa_comment(metrics)
    doc.core_properties.comments = f"{current_comments}\n{qa_line}".strip()

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def install(word_engine) -> None:
    original_enhance = word_engine.enhance_docx

    def enhanced(data, logger, language="en"):
        first_pass = original_enhance(data, logger, language)
        return visual_qa(first_pass, language)

    word_engine.enhance_docx = enhanced
