from __future__ import annotations

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

from curriculum_models import LongPlan, MediumPlan
from curriculum_ai import localize_grade, localize_subject, normalize_language

ROOT = Path(__file__).resolve().parent
MEDIUM_TEMPLATE = ROOT / "word_templates" / "medium_term_template.docx"
LONG_TEMPLATE = ROOT / "word_templates" / "long_term_template.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS = {"w": W_NS, "w15": W15_NS, "wp": WP_NS, "a": A_NS}


def _is_arabic(language: str) -> bool:
    return str(language or "").strip().casefold() in {"arabic", "ar", "العربية", "عربي"}


def _set_bidi(paragraph, enabled: bool) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if enabled and bidi is None:
        ppr.append(OxmlElement("w:bidi"))
    elif not enabled and bidi is not None:
        ppr.remove(bidi)


def _clear_cell(cell) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)


def _set_cell_text(cell, text: str, size: float = 8, bold_first: bool = False, rtl: bool = False, center: bool = False) -> None:
    _clear_cell(cell)
    for idx, line in enumerate((text or "").splitlines() or [""]):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else (WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        _set_bidi(p, rtl)
        run = p.add_run(line)
        run.bold = bold_first and idx == 0
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _set_exact_height(row, twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for old in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(old)
    node = OxmlElement("w:trHeight")
    node.set(qn("w:val"), str(twips))
    node.set(qn("w:hRule"), "exact")
    tr_pr.append(node)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _compact_cell(cell, size: float, rtl: bool, bold: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 30), ("start", 40), ("bottom", 30), ("end", 40)):
        item = margins.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            margins.append(item)
        item.set(qn("w:w"), str(value)); item.set(qn("w:type"), "dxa")
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = .85
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
        _set_bidi(p, rtl)
        for run in p.runs:
            run.font.size = Pt(size); run.font.name = "Arial"; run.bold = bold or run.bold
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")




def _remove_trailing_empty_paragraphs(doc) -> None:
    body = doc._element.body
    children = list(body)
    last_table = max((i for i, el in enumerate(children) if el.tag == qn("w:tbl")), default=-1)
    for i, el in enumerate(list(body)):
        if i > last_table and el.tag == qn("w:p") and not "".join(el.itertext()).strip():
            body.remove(el)


LONG_EN = {
    "Area": ["Select...", "Governance", "Teacher Capacity", "Curriculum Integration", "Infrastructure", "Student Exposure", "Monitoring and Evaluation"],
    "Milestone_Target": ["Select...", "AI Lead Coordinator appointed and implementation plan submitted", "All staff complete AI Foundations and Safety Training", "All staff complete AI Safety and Ethics Training", "Approved AI tools and LMS integration deployed", "AI embedded in at least one core subject per grade level", "Termly review of AI curriculum integration and student exposure", "Student AI literacy evidence collected and reviewed", "Parental consent and data-sharing processes implemented", "School Readiness Self-Assessment completed", "AI Implementation Plan submitted to ADEK"],
    "Responsible_Person": ["Select...", "Principal", "Vice Principal or Deputy Head", "AI Lead Coordinator", "Head of Department", "IT or Digital Lead", "Classroom Teacher", "School Leadership Team", "All Teaching Staff"],
    "Target_Date": ["Select...", "August 2026", "September 2026", "October 2026", "November 2026", "December 2026", "January 2027", "February 2027", "March 2027", "April 2027", "May 2027", "June 2027", "July 2027", "End of Term 1", "End of Term 2", "End of Term 3", "Academic Year 2026-2027"],
    "Status": ["Select...", "Not Started", "Planning in Progress", "In Progress", "Completed", "Delayed", "At Risk", "Requires Review"],
}
LONG_AR = {
    "Area": ["اختر...", "الحوكمة", "بناء قدرات المعلمين", "تكامل المنهج", "البنية التحتية", "تعرض الطلبة", "المتابعة والتقويم"],
    "Milestone_Target": ["اختر...", "تعيين منسق قيادة الذكاء الاصطناعي واعتماد خطة التنفيذ", "إكمال جميع الموظفين تدريب أساسيات الذكاء الاصطناعي والسلامة", "إكمال جميع الموظفين تدريب السلامة والأخلاقيات", "تفعيل الأدوات المعتمدة وربطها بمنصة إدارة التعلم", "دمج الذكاء الاصطناعي في مادة أساسية واحدة على الأقل لكل صف", "تنفيذ مراجعة فصلية لتكامل المنهج وتعرض الطلبة", "جمع أدلة ثقافة الذكاء الاصطناعي لدى الطلبة ومراجعتها", "تطبيق إجراءات موافقة أولياء الأمور ومشاركة البيانات", "إكمال التقييم الذاتي لجاهزية المدرسة", "تقديم خطة تطبيق الذكاء الاصطناعي إلى دائرة التعليم والمعرفة"],
    "Responsible_Person": ["اختر...", "مدير المدرسة", "نائب المدير", "منسق قيادة الذكاء الاصطناعي", "رئيس القسم", "مسؤول التقنية والتحول الرقمي", "معلم الصف", "فريق القيادة المدرسية", "جميع أعضاء الهيئة التدريسية"],
    "Target_Date": ["اختر...", "أغسطس 2026", "سبتمبر 2026", "أكتوبر 2026", "نوفمبر 2026", "ديسمبر 2026", "يناير 2027", "فبراير 2027", "مارس 2027", "أبريل 2027", "مايو 2027", "يونيو 2027", "يوليو 2027", "نهاية الفصل الدراسي الأول", "نهاية الفصل الدراسي الثاني", "نهاية الفصل الدراسي الثالث", "العام الأكاديمي 2026-2027"],
    "Status": ["اختر...", "لم يبدأ", "التخطيط جارٍ", "قيد التنفيذ", "مكتمل", "متأخر", "معرض للخطر", "يحتاج إلى مراجعة"],
}
MEDIUM_EN = {"AI Integration Approach": ["Choose an item...", "AI-supported teaching for planning and resources", "AI literacy learning for students", "Unplugged AI activity"], "Guardrails": ["Choose an item...", "Approved tool with cultural and safety filters", "Strict teacher-provided prompt template", "Closed-environment tool without external web access", "Real-time teacher monitoring"], "Cognitive Integrity": ["Choose an item...", "Students submit a draft before AI support", "Students explain reasoning behind AI-supported outputs", "Students evaluate and correct AI errors", "AI used only after an independent attempt"], "Safeguarding": ["Choose an item...", "Anonymous accounts and no personal data shared", "Parental consent obtained", "Explicit teaching about bias, hallucinations, and integrity", "All of the above"], "Area": LONG_EN["Area"], "Milestone": ["Choose an item...", "Appoint AI Lead Coordinator", "Complete AI Foundations Training", "Complete AI Safety and Ethics Training", "Deploy Approved AI Tools", "Integrate AI in one core subject", "Conduct Termly Review"], "Responsible Person": ["Choose an item...", "Principal", "Vice Principal", "AI Lead Coordinator", "Head of Department", "IT or Digital Lead", "Classroom Teacher"], "Status": ["Choose an item...", "Not Started", "In Progress", "Completed", "Delayed", "At Risk"]}
MEDIUM_AR = {"AI Integration Approach": ["اختر...", "تدريس مدعوم بالذكاء الاصطناعي للتخطيط والمصادر", "تعلم ثقافة الذكاء الاصطناعي للطلبة", "نشاط غير رقمي للذكاء الاصطناعي"], "Guardrails": ["اختر...", "أداة معتمدة مزودة بضوابط ثقافية وأمنية", "نموذج أوامر صارم يقدمه المعلم", "أداة مغلقة دون وصول خارجي للإنترنت", "متابعة مباشرة من المعلم"], "Cognitive Integrity": ["اختر...", "يقدم الطلبة مسودة قبل الاستعانة بالأداة", "يشرح الطلبة منطق المخرجات المدعومة رقمياً", "يقيم الطلبة الأخطاء الرقمية ويصححونها", "تستخدم الأداة بعد محاولة مستقلة فقط"], "Safeguarding": ["اختر...", "حسابات مجهولة وعدم مشاركة بيانات شخصية", "الحصول على موافقة ولي الأمر", "تعليم صريح حول التحيز والهلوسة والنزاهة", "جميع ما سبق"], "Area": LONG_AR["Area"], "Milestone": ["اختر...", "تعيين منسق قيادة الذكاء الاصطناعي", "إكمال تدريب الأساسيات", "إكمال تدريب السلامة والأخلاقيات", "تفعيل الأدوات المعتمدة", "دمج الذكاء الاصطناعي في مادة أساسية", "إجراء مراجعة فصلية"], "Responsible Person": ["اختر...", "مدير المدرسة", "نائب المدير", "منسق قيادة الذكاء الاصطناعي", "رئيس القسم", "مسؤول التقنية", "معلم الصف"], "Status": ["اختر...", "لم يبدأ", "قيد التنفيذ", "مكتمل", "متأخر", "معرض للخطر"]}

AR_LABELS = {
    "Long Term Plan": "الخطة طويلة المدى", "AI Implementation & Compliance Tracking (School Level)": "تتبع تطبيق الذكاء الاصطناعي والامتثال على مستوى المدرسة",
    "Area": "المجال", "Milestone / Target": "الإنجاز / الهدف", "Responsible Person": "المسؤول", "Target Date": "التاريخ المستهدف", "Status": "الحالة",
    "Week": "الأسبوع", "Content": "المحتوى", "Learning Objectives": "أهداف التعلم", "AI Literacy Objective / Integration": "هدف ثقافة الذكاء الاصطناعي / التكامل", "Resources& AI Tools": "المصادر وأدوات الذكاء الاصطناعي",
    "MID-TERM BREAK": "إجازة منتصف الفصل", "Summative Assessments": "التقويمات الختامية",
    "Link to National Identity Mark:": "الارتباط بعلامة الهوية الوطنية:", "Domain:": "المجال:", "Dimensions:": "الأبعاد:",
    "CULTURE": "الثقافة", "VALUES": "القيم", "CITIZENSHIP": "المواطنة", "Arabic language": "اللغة العربية",
    "Respect": "الاحترام", "Belonging": "الانتماء", "History": "التاريخ", "Compassion": "التعاطف",
    "Volunteering": "التطوع", "Heritage": "التراث", "Global Understanding": "الفهم العالمي", "Conservation": "المحافظة على الموارد",
    "1. AI Integration Approach": "1. نهج دمج الذكاء الاصطناعي", "2. Guardrails & Prompt Controls": "2. الضوابط والتحكم في الأوامر",
    "3. Cognitive Integrity Strategy": "3. استراتيجية النزاهة المعرفية", "4. AI Safeguarding & Responsible Use": "4. السلامة والاستخدام المسؤول للذكاء الاصطناعي",
    "5. AI Implementation & Compliance Tracking": "5. تتبع التطبيق والامتثال",
    "How is the AI tool restricted to ensure cultural alignment, safety, and on-topic focus?": "كيف تُقيد أداة الذكاء الاصطناعي لضمان التوافق الثقافي والسلامة والتركيز على المهمة؟",
    "How will you ensure students do not use AI for cognitive offloading?": "كيف ستضمن عدم اعتماد الطلبة على الذكاء الاصطناعي بدلاً من التفكير المستقل؟",
    "Detail guardrails, data privacy measures, and ethical considerations for this term.": "وضح الضوابط وإجراءات خصوصية البيانات والاعتبارات الأخلاقية لهذا الفصل.",
    "Select Date": "اختر التاريخ", "AI Integration Approach": "نهج دمج الذكاء الاصطناعي",
    "Guardrails & Prompt Controls": "الضوابط والتحكم في الأوامر", "Cognitive Integrity Strategy": "استراتيجية النزاهة المعرفية",
    "AI Safeguarding & Responsible Use": "السلامة والاستخدام المسؤول للذكاء الاصطناعي",
    "AI Implementation & Compliance Tracking": "تتبع التطبيق والامتثال",
    "Domain": "المجال", "Dimensions": "الأبعاد", "Arabic": "اللغة العربية", "language": "",
    "Global": "الفهم", "Understanding": "العالمي",
}


def _control_key(pr) -> str:
    for name in ("tag", "alias"):
        node = pr.find(f"w:{name}", namespaces=NS)
        if node is not None and node.get(f"{{{W_NS}}}val"):
            return node.get(f"{{{W_NS}}}val")
    return ""


def _set_control_text(sdt, text: str) -> None:
    content = sdt.find("w:sdtContent", namespaces=NS)
    if content is None:
        content = etree.SubElement(sdt, f"{{{W_NS}}}sdtContent")
    texts = content.xpath(".//w:t", namespaces=NS)
    if texts:
        texts[0].text = text
        for node in texts[1:]: node.text = ""
    else:
        run = etree.SubElement(content, f"{{{W_NS}}}r"); node = etree.SubElement(run, f"{{{W_NS}}}t"); node.text = text


def _patch_dropdowns(root, arabic: bool, kind: str) -> None:
    options = (LONG_AR if arabic else LONG_EN) if kind == "long" else (MEDIUM_AR if arabic else MEDIUM_EN)
    for sdt in root.xpath(".//w:sdt[w:sdtPr/w:dropDownList]", namespaces=NS):
        pr = sdt.find("w:sdtPr", namespaces=NS); dropdown = pr.find("w:dropDownList", namespaces=NS); key = _control_key(pr)
        values = options.get(key)
        if not values: continue
        current = "".join(sdt.xpath(".//w:sdtContent//w:t/text()", namespaces=NS)).casefold()
        selected = values[0]
        if key == "Area" and ("curriculum integration" in current or "تكامل المنهج" in current): selected = "تكامل المنهج" if arabic else "Curriculum Integration"
        elif key == "Area" and ("student exposure" in current or "تعرض الطلبة" in current): selected = "تعرض الطلبة" if arabic else "Student Exposure"
        for child in list(dropdown): dropdown.remove(child)
        for value in values:
            item = etree.SubElement(dropdown, f"{{{W_NS}}}listItem"); item.set(f"{{{W_NS}}}displayText", value); item.set(f"{{{W_NS}}}value", "" if value == values[0] else value)
        for lock in pr.findall("w:lock", namespaces=NS): pr.remove(lock)
        for name in ("appearance", "color"):
            for old in pr.findall(f"w15:{name}", namespaces=NS): pr.remove(old)
        appearance = etree.SubElement(pr, f"{{{W15_NS}}}appearance"); appearance.set(f"{{{W15_NS}}}val", "boundingBox")
        color = etree.SubElement(pr, f"{{{W15_NS}}}color"); color.set(f"{{{W15_NS}}}val", "4472C4")
        _set_control_text(sdt, selected)


def _postprocess(path: Path, language: str, kind: str) -> None:
    arabic = _is_arabic(language)
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(path.parent)); os.close(fd); tmp = Path(tmp_name)
    try:
        with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename == "word/document.xml":
                    root = etree.fromstring(data)
                    if arabic:
                        for node in root.xpath(".//w:t", namespaces=NS):
                            text = node.text or ""
                            for en, ar in sorted(AR_LABELS.items(), key=lambda x: len(x[0]), reverse=True): text = text.replace(en, ar)
                            node.text = text
                    _patch_dropdowns(root, arabic, kind)
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                elif info.filename == "word/settings.xml":
                    root = etree.fromstring(data)
                    for node in root.xpath(".//w:documentProtection", namespaces=NS): node.getparent().remove(node)
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                elif info.filename.startswith("word/header") and info.filename.endswith(".xml"):
                    root = etree.fromstring(data)
                    for node in root.xpath(".//wp:extent | .//a:xfrm/a:ext", namespaces=NS):
                        for attr in ("cx", "cy"):
                            value = node.get(attr)
                            if value and value.isdigit(): node.set(attr, str(int(int(value) * .75)))
                    data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                zout.writestr(info, data)
        tmp.replace(path)
    finally:
        tmp.unlink(missing_ok=True)


def _localized_meta(meta: dict) -> tuple[bool, str, str, str, str]:
    language = normalize_language(meta.get("language", "English"))
    rtl = language == "Arabic"
    subject = localize_subject(meta.get("subject", ""), language)
    grade = localize_grade(meta.get("grade", ""), language)
    teacher = str(meta.get("teacher", "")).strip()
    academic_year = str(meta.get("academic_year", "2026-2027")).strip()
    return rtl, subject, grade, teacher, academic_year


def build_medium(meta: dict, plan: MediumPlan, output_path: Path) -> Path:
    doc = Document(MEDIUM_TEMPLATE); rtl, subject, grade, teacher, academic_year = _localized_meta(meta)
    if doc.paragraphs:
        p = doc.paragraphs[0]; p.text = plan.title; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs: run.bold = True; run.font.size = Pt(15)
    top = doc.tables[0]
    if rtl:
        _set_cell_text(top.cell(0, 0), f"العام: \u200e{academic_year}\u200e\nالمعلم: {teacher}", 9, True, True)
        _set_cell_text(top.cell(0, 1), f"الفصل الدراسي: الأول\nالصف: {grade}", 9, True, True)
        _set_cell_text(top.cell(0, 2), f"الأهداف العامة:\n{plan.targets}", 8, True, True)
        _set_cell_text(top.cell(1, 0), f"العنوان: {plan.title}", 9, True, True)
        _set_cell_text(top.cell(2, 0), "المدة: 31 أغسطس - 11 ديسمبر 2026 (14 أسبوعاً)", 9, True, True)
    else:
        _set_cell_text(top.cell(0, 0), f"Year: {academic_year}\nTeacher: {teacher}", 9, True)
        _set_cell_text(top.cell(0, 1), f"Term: 1\nGrade: {grade}", 9, True)
        _set_cell_text(top.cell(0, 2), f"Targets:\n{plan.targets}", 8, True)
        _set_cell_text(top.cell(1, 0), f"Title: {plan.title}", 9, True)
        _set_cell_text(top.cell(2, 0), "Duration: 31 Aug - 11 Dec 2026 (14 weeks)", 9, True)
    weekly = doc.tables[1]; rows = [1,2,3,4,5,6,8,9,10,11,12,13,14,15]
    if rtl:
        headers = ["الأسبوع", "المحتوى", "أهداف التعلم", "هدف ثقافة الذكاء الاصطناعي / التكامل", "المصادر وأدوات الذكاء الاصطناعي"]
        dates = ["31 أغسطس - 4 سبتمبر", "7 - 11 سبتمبر", "14 - 18 سبتمبر", "21 - 25 سبتمبر", "28 سبتمبر - 2 أكتوبر", "5 - 9 أكتوبر", "19 - 23 أكتوبر", "26 - 30 أكتوبر", "2 - 6 نوفمبر", "9 - 13 نوفمبر", "16 - 20 نوفمبر", "23 - 27 نوفمبر", "30 نوفمبر - 4 ديسمبر", "7 - 11 ديسمبر"]
    else:
        headers = ["Week", "Content", "Learning Objectives", "AI Literacy Objective / Integration", "Resources and AI Tools"]
        dates = ["31 Aug - 4 Sep", "7 - 11 Sep", "14 - 18 Sep", "21 - 25 Sep", "28 Sep - 2 Oct", "5 - 9 Oct", "19 - 23 Oct", "26 - 30 Oct", "2 - 6 Nov", "9 - 13 Nov", "16 - 20 Nov", "23 - 27 Nov", "30 Nov - 4 Dec", "7 - 11 Dec"]
    for col, label in enumerate(headers): _set_cell_text(weekly.cell(0,col), label, 7.4, True, rtl, True)
    for item, row_idx, week_no, date_text in zip(plan.weeks[:14], rows, range(1,15), dates):
        _set_cell_text(weekly.rows[row_idx].cells[0], f"{'الأسبوع' if rtl else 'W'} {week_no}\n{date_text}", 7.0, True, rtl)
        row = weekly.rows[row_idx]; _cant_split(row)
        _set_cell_text(row.cells[1], item.content, 6.7, rtl=rtl); _set_cell_text(row.cells[2], item.learning_objectives, 6.4, rtl=rtl)
        _set_cell_text(row.cells[3], item.ai_literacy, 6.2, rtl=rtl); _set_cell_text(row.cells[4], item.resources, 6.2, rtl=rtl)
    _set_cell_text(weekly.rows[7].cells[0], "إجازة منتصف الفصل\n12 - 18 أكتوبر 2026\nاستئناف الدراسة 19 أكتوبر" if rtl else "MID-TERM BREAK\n12 - 18 Oct 2026\nClasses resume 19 Oct", 6.8, True, rtl)
    for col in range(1,5): _set_cell_text(weekly.rows[7].cells[col], "", 7, rtl=rtl)
    summary = doc.tables[2]
    labels = ("فرص التقويم", "مهارات القرن الحادي والعشرين", "المفردات والكلمات المفتاحية") if rtl else ("Assessment Opportunities", "21st Century Skills", "Vocabulary and Key Words")
    values = (plan.assessment_opportunities, plan.century_skills, plan.vocabulary)
    for col in range(3): _set_cell_text(summary.cell(0,col), f"{labels[col]}:\n{values[col]}", 6.8, True, rtl)
    footer = doc.tables[3]
    detail_labels = ("الارتباط بالتوجه العام للمدرسة والكفاءات والقيم", "المواطنة العالمية", "الروابط الأفقية بين المواد") if rtl else ("EPS Guiding Statement, Competences, HQL and Values", "Global Citizenship", "Cross-curricular and Horizontal Articulation")
    detail_values = (plan.eps_guiding_statement, plan.global_citizenship, plan.cross_curricular)
    for row in range(3): _set_cell_text(footer.cell(row,0), f"{detail_labels[row]}:\n{detail_values[row]}", 7, True, rtl)
    output_path.parent.mkdir(parents=True, exist_ok=True); doc.save(output_path); _postprocess(output_path, meta.get("language", "English"), "medium"); return output_path


def build_long(meta: dict, plan: LongPlan, output_path: Path) -> Path:
    doc = Document(LONG_TEMPLATE)
    for section in doc.sections: section.top_margin = Inches(.90); section.header_distance = Inches(0)
    rtl, subject, grade, teacher, academic_year = _localized_meta(meta); metadata = doc.tables[0]
    if rtl:
        _set_cell_text(metadata.cell(0,0), f"المادة: {subject}\nالمعلم: {teacher}", 9, True, True)
        _set_cell_text(metadata.cell(0,1), f"الصف: {grade}", 9, True, True)
        _set_cell_text(metadata.cell(0,2), f"العام الأكاديمي: \u200e{academic_year}\u200e", 9, True, True)
    else:
        _set_cell_text(metadata.cell(0,0), f"Subject: {subject}\nTeacher: {teacher}", 9, True)
        _set_cell_text(metadata.cell(0,1), f"Grade Group: {grade}", 9, True)
        _set_cell_text(metadata.cell(0,2), f"Academic Year: {academic_year}", 9, True)
    curriculum = doc.tables[1]
    if rtl:
        headers = [(0,1,"الفصل الدراسي الأول\nمن 31 أغسطس 2026 إلى 11 ديسمبر 2026 (14 أسبوعاً)"),(0,3,"الفصل الدراسي الثاني\nمن 4 يناير 2027 إلى 2 أبريل 2027 (13 أسبوعاً)"),(0,5,"الفصل الدراسي الثالث\nمن 12 أبريل 2027 إلى 2 يوليو 2027 (12 أسبوعاً)")]
        for r,c,text in headers: _set_cell_text(curriculum.cell(r,c), text, 8.1, True, True, True)
        half_headers = ["الفترة الأولى - الفصل الأول\nمن 31 أغسطس إلى 9 أكتوبر (6 أسابيع)", "الفترة الثانية - الفصل الأول\nمن 19 أكتوبر إلى 11 ديسمبر (8 أسابيع)", "الفترة الأولى - الفصل الثاني\nمن 4 يناير إلى 19 فبراير (7 أسابيع)", "الفترة الثانية - الفصل الثاني\nمن 22 فبراير إلى 2 أبريل (6 أسابيع)", "الفترة الأولى - الفصل الثالث\nمن 12 أبريل إلى 21 مايو (6 أسابيع)", "الفترة الثانية - الفصل الثالث\nمن 24 مايو إلى 2 يوليو (6 أسابيع)"]
        for col,text in enumerate(half_headers,1): _set_cell_text(curriculum.cell(1,col), text, 7.3, True, True, True)
        _set_cell_text(curriculum.cell(3,0), "التقويمات\nالختامية", 7.2, True, True, True)
    for idx, half in enumerate(plan.half_terms[:6],1):
        _set_cell_text(curriculum.cell(2,idx), half.content, 5.55 if rtl else 6.0, rtl=rtl); _set_cell_text(curriculum.cell(3,idx), half.summative_assessment, 5.15 if rtl else 5.7, rtl=rtl)
    _cant_split(curriculum.rows[2]); _cant_split(curriculum.rows[3])
    compliance = doc.tables[2]
    for row_idx,row in enumerate(compliance.rows[:5]):
        _cant_split(row)
        if row_idx: _set_exact_height(row, 300)
        for cell in row.cells: _compact_cell(cell, 5.9 if row_idx else 6.4, rtl, row_idx == 0)
    _remove_trailing_empty_paragraphs(doc); output_path.parent.mkdir(parents=True, exist_ok=True); doc.save(output_path); _postprocess(output_path, meta.get("language", "English"), "long"); return output_path
