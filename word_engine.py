from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from math_format import build_equation

RLM = "\u200f"
EQ_RE = re.compile(r"\[\[EQ:(.+?)\]\]")
NUMBERED_RE = re.compile(r"^\s*[\u200e\u200f]*(\d+)[\.)-]?\s*(.*)$")
INLINE_NUMBER_RE = re.compile(r"([^\n])\s+((?:[\u200e\u200f])?\d{1,2}[\.)]\s+(?=[^\d]))")
BRAND_MARKERS = ("www.magdymath.com", "prepared by mr. magdy elsayed")

NAVY_HEX = "0A3568"
BLUE_HEX = "2E75B6"
LIGHT_BLUE_HEX = "D9EAF7"
PALE_BLUE_HEX = "EDF5FB"
WHITE_HEX = "FFFFFF"
BORDER_HEX = "7FA6C9"
NAVY = RGBColor(10, 53, 104)
BLUE = RGBColor(31, 78, 121)
WHITE = RGBColor(255, 255, 255)
DARK = RGBColor(28, 40, 55)

HEADING_PREFIXES = (
    "تمهيد", "سؤال تشخيصي", "الاستجابة المتوقعة", "مثال محلول", "تدريب موجه",
    "تطبيق فردي", "نشاط تعاوني", "سؤال إثرائي", "سؤال تفكير عليا", "بطاقة خروج",
    "دور المعلم", "دور الطلاب", "دعم", "المستوى المتوقع", "متقدمون", "تكييفات",
    "خطأ متوقع", "معيار النجاح", "المنتج المتوقع", "المهمة", "التوجيه",
    "hook", "diagnostic question", "expected response", "worked example", "guided practice",
    "independent practice", "independent application", "collaborative task", "exit ticket",
    "teacher role", "student role", "support", "expected level", "advanced", "misconception",
    "success measure", "learning evidence", "hots",
)

ARABIC_LABELS = {
    "lesson plan ay2026-2027": "خطة الدرس للعام الأكاديمي 2026–2027",
    "lesson plan": "خطة الدرس",
    "teacher:": "المعلم:",
    "subject:": "المادة:",
    "date:": "التاريخ:",
    "class:": "الصف:",
    "topic:": "موضوع الدرس:",
    "periods:": "عدد الحصص:",
    "key words:": "الكلمات المفتاحية:",
    "key words": "الكلمات المفتاحية",
    "keywords:": "الكلمات المفتاحية:",
    "primary sdg focus:": "الهدف الرئيس للتنمية المستدامة:",
    "strategies:": "استراتيجيات التدريس:",
    "learning outcomes:": "نواتج التعلم:",
    "success criteria:": "معايير النجاح:",
    "lesson structure": "هيكل الدرس",
    "starter (10 mins)": "التمهيد (10 دقائق)",
    "main activities (15 mins)": "الأنشطة الرئيسة (15 دقيقة)",
    "teacher-led": "دور المعلم",
    "student-led": "دور الطلاب",
    "plenary (10 mins)": "الخاتمة (10 دقائق)",
    "kpi afl assignment task:": "مهمة التقويم التكويني KPI/AFL:",
    "21 century skills prompts:": "مهارات القرن الحادي والعشرين:",
    "resources: digital technology/ software teacher student": "المصادر والتقنيات الرقمية:",
    "my identity: links to uae culture": "هويتي: روابط بالثقافة الإماراتية:",
    "student competency framework: links to values and competencies (from sdg competency program)": "إطار كفاءات الطالب: القيم والكفاءات:",
}

SECTION_TITLES = {
    "lesson plan", "خطة الدرس", "lesson structure", "هيكل الدرس",
    "lesson plan ay2026-2027", "خطة الدرس للعام الأكاديمي 2026–2027",
}

LABEL_KEYS = set(ARABIC_LABELS.keys()) | set(ARABIC_LABELS.values())


def _normalise_label(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip()).casefold()


def _normalize_lines(text, clean_text):
    value = clean_text(text)
    value = INLINE_NUMBER_RE.sub(r"\1\n\2", value)
    value = re.sub(r"\n[ \t]+", "\n", value)
    return [line.strip() for line in value.split("\n") if line.strip()] or [""]


def _font_name(rtl: bool) -> str:
    return "Arial" if rtl else "Aptos"


def _shade(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _set_cell_border(cell, color: str = BORDER_HEX, size: str = "6") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def _set_paragraph_accent(paragraph, rtl: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge_name = "right" if rtl else "left"
    edge = borders.find(qn(f"w:{edge_name}"))
    if edge is None:
        edge = OxmlElement(f"w:{edge_name}")
        borders.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "16")
    edge.set(qn("w:space"), "5")
    edge.set(qn("w:color"), BLUE_HEX)


def _add_text_run(paragraph, text, app_module, rtl, size, bold=False, color=None):
    if not text:
        return
    run = paragraph.add_run(text)
    app_module.set_run_font(run, _font_name(rtl), size, bold)
    if color is not None:
        run.font.color.rgb = color


def _heading_split(line: str):
    if ":" not in line:
        return None
    possible, rest = line.split(":", 1)
    key = possible.strip().casefold()
    group_heading = bool(re.match(r"^(?:مجموعة|المجموعة|group)\s*[0-9٠-٩\w]+", key))
    if len(possible.strip()) <= 46 and (key in HEADING_PREFIXES or group_heading):
        return possible.strip(), rest.lstrip()
    return None


def _append_mixed_content(paragraph, line, app_module, rtl, size):
    numbered = NUMBERED_RE.match(line)
    if numbered:
        number_text = numbered.group(1)
        body = numbered.group(2)
        prefix = f"{RLM}{number_text}. {RLM}" if rtl else f"{number_text}. "
        _add_text_run(paragraph, prefix, app_module, rtl, size, True, BLUE)
        line = body

    heading_match = _heading_split(line)
    if heading_match:
        _set_paragraph_accent(paragraph, rtl)
        paragraph.paragraph_format.keep_with_next = True
        _add_text_run(paragraph, heading_match[0] + ": ", app_module, rtl, size + 0.35, True, NAVY)
        line = heading_match[1]

    position = 0
    for equation in EQ_RE.finditer(line):
        before = line[position:equation.start()]
        _add_text_run(paragraph, before, app_module, rtl, size, False, DARK)
        paragraph._element.append(build_equation(equation.group(1).strip()))
        position = equation.end()
    _add_text_run(paragraph, line[position:], app_module, rtl, size, False, DARK)


def install_word_upgrade(app_module):
    def set_cell_text(cell, text, lang="en", size=8.0, bold=False):
        rtl = lang == "ar"
        body_size = max(float(size) + 1.0, 8.6)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        lines = _normalize_lines(text, app_module.clean_text)

        for index, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            equation_only_match = EQ_RE.fullmatch(line)
            equation_only = bool(equation_only_match)
            # In table cells, side alignment keeps equations visually connected to the instruction.
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            app_module.set_paragraph_bidi(paragraph, rtl and not equation_only)
            paragraph.paragraph_format.space_before = Pt(1.0 if equation_only else 0)
            paragraph.paragraph_format.space_after = Pt(1.0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.paragraph_format.keep_together = True
            if rtl:
                paragraph.paragraph_format.right_indent = Pt(2)
            else:
                paragraph.paragraph_format.left_indent = Pt(2)

            if equation_only:
                paragraph._element.append(build_equation(equation_only_match.group(1).strip()))
            else:
                _append_mixed_content(paragraph, line, app_module, rtl, body_size)

    app_module.set_cell_text = set_cell_text
    original = app_module.generate_docx

    def generate_docx(lesson):
        return enhance_docx(original(lesson), app_module.logger, lesson.language)

    app_module.generate_docx = generate_docx


def remove_fixed_height(row):
    props = row._tr.get_or_add_trPr()
    for node in list(props.findall(qn("w:trHeight"))):
        props.remove(node)


def clear_repeat_header(row):
    props = row._tr.get_or_add_trPr()
    for node in list(props.findall(qn("w:tblHeader"))):
        props.remove(node)


def _set_cant_split(row, enabled: bool) -> None:
    props = row._tr.get_or_add_trPr()
    existing = props.find(qn("w:cantSplit"))
    if enabled and existing is None:
        props.append(OxmlElement("w:cantSplit"))
    elif not enabled and existing is not None:
        props.remove(existing)


def set_cell_margins(cell, top=75, start=100, bottom=75, end=100):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_branding_line(doc):
    for section in doc.sections:
        for paragraph in list(section.header.paragraphs):
            text = paragraph.text.strip().lower()
            if any(marker in text for marker in BRAND_MARKERS):
                _remove_paragraph(paragraph)
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip().lower()
        if any(marker in text for marker in BRAND_MARKERS):
            _remove_paragraph(paragraph)


def _table_signature(table):
    labels = []
    seen_cells = set()
    for row in table.rows:
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen_cells:
                continue
            seen_cells.add(key)
            text = _normalise_label(cell.text)
            if text in LABEL_KEYS or any(label in text for label in SECTION_TITLES):
                labels.append(text)
    return len(table.rows), len(table.columns), tuple(sorted(set(labels)))


def remove_duplicate_tables(doc):
    seen = set()
    for table in list(doc.tables):
        signature = _table_signature(table)
        if len(signature[2]) < 4:
            continue
        if signature in seen:
            parent = table._element.getparent()
            if parent is not None:
                parent.remove(table._element)
        else:
            seen.add(signature)


def _replace_paragraph_text(paragraph, text: str, rtl: bool, bold: bool = False, color=None, size=9.0):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = _font_name(rtl)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.append(fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(key), _font_name(rtl))
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def translate_template_labels(doc, language: str):
    if language != "ar":
        return
    for paragraph in doc.paragraphs:
        key = _normalise_label(paragraph.text)
        if key in ARABIC_LABELS:
            _replace_paragraph_text(paragraph, ARABIC_LABELS[key], True, True, NAVY, 10)
    for table in doc.tables:
        seen_cells = set()
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                key = _normalise_label(cell.text)
                if key in ARABIC_LABELS:
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    _replace_paragraph_text(paragraph, ARABIC_LABELS[key], True, True, WHITE, 8.8)


def _is_label_cell(cell) -> bool:
    key = _normalise_label(cell.text)
    return key in LABEL_KEYS


def _is_section_cell(cell) -> bool:
    key = _normalise_label(cell.text)
    return key in SECTION_TITLES


def _style_label_cell(cell, rtl: bool, section: bool = False) -> None:
    _shade(cell._tc, NAVY_HEX if section else BLUE_HEX)
    _set_cell_border(cell, NAVY_HEX if section else BORDER_HEX, "7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if section else (WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            run.font.name = _font_name(rtl)
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10.5 if section else 8.8)


def style_document_tables(doc, language: str):
    rtl = language == "ar"
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            clear_repeat_header(row)
            remove_fixed_height(row)
            row_text_length = sum(len(cell.text) for cell in row.cells)
            _set_cant_split(row, row_text_length < 900)
            seen_cells = set()
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                set_cell_margins(cell)
                _set_cell_border(cell)
                section = _is_section_cell(cell)
                if section or _is_label_cell(cell):
                    _style_label_cell(cell, rtl, section)
                    continue
                _shade(cell._tc, WHITE_HEX)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(1.0)
                    paragraph.paragraph_format.line_spacing = 1.05
                    if rtl:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        if run.font.size is None or run.font.size.pt < 8.4:
                            run.font.size = Pt(8.4)
                        run.font.name = _font_name(rtl)


def _clean_empty_paragraphs(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = list(cell.paragraphs)
                while len(paragraphs) > 1 and not paragraphs[-1].text.strip() and len(paragraphs[-1]._p) <= 1:
                    _remove_paragraph(paragraphs[-1])
                    paragraphs = list(cell.paragraphs)


def _set_page_layout(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.48)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.58)
        section.right_margin = Inches(0.58)
        section.header_distance = Inches(0.08)
        section.footer_distance = Inches(0.18)


def enhance_docx(data, logger, language="en"):
    doc = Document(io.BytesIO(data))
    remove_branding_line(doc)
    remove_duplicate_tables(doc)
    translate_template_labels(doc, language)
    style_document_tables(doc, language)
    _clean_empty_paragraphs(doc)
    _set_page_layout(doc)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
