from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(10, 53, 104)
MID_BLUE = RGBColor(46, 117, 182)
GREY = RGBColor(92, 108, 125)


def _field(paragraph, instruction: str, fallback: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_text, separate, result, end])


def _set_bidi(paragraph, enabled: bool) -> None:
    properties = paragraph._p.get_or_add_pPr()
    bidi = properties.find(qn("w:bidi"))
    if enabled and bidi is None:
        properties.append(OxmlElement("w:bidi"))
    elif not enabled and bidi is not None:
        properties.remove(bidi)


def _add_footer_border(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    top = borders.find(qn("w:top"))
    if top is None:
        top = OxmlElement("w:top")
        borders.append(top)
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "5")
    top.set(qn("w:color"), "2E75B6")


def _style_run(run, size=8.0, bold=False, color=GREY, font="Arial") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.append(fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(key), font)


def _professional_footer(doc, language: str) -> None:
    rtl = language == "ar"
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(0)
        _set_bidi(paragraph, rtl)
        _add_footer_border(paragraph)

        site = paragraph.add_run("WWW.MAGDYMATH.COM")
        _style_run(site, 8.2, True, NAVY)
        separator = paragraph.add_run("   |   ")
        _style_run(separator, 8.0, False, GREY)
        label = paragraph.add_run("صفحة " if rtl else "Page ")
        _style_run(label, 8.0, False, GREY)
        _field(paragraph, "PAGE")
        middle = paragraph.add_run(" من " if rtl else " of ")
        _style_run(middle, 8.0, False, GREY)
        _field(paragraph, "NUMPAGES")


def _set_cell_width(cell, inches: float) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(int(inches * 1440)))
    width.set(qn("w:type"), "dxa")


def _set_row_keep(row, enabled: bool) -> None:
    properties = row._tr.get_or_add_trPr()
    node = properties.find(qn("w:cantSplit"))
    if enabled and node is None:
        properties.append(OxmlElement("w:cantSplit"))
    elif not enabled and node is not None:
        properties.remove(node)


def _paragraph_controls(paragraph, keep_with_next=False) -> None:
    paragraph.paragraph_format.widow_control = True
    paragraph.paragraph_format.keep_together = True
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True
    properties = paragraph._p.get_or_add_pPr()
    if properties.find(qn("w:suppressAutoHyphens")) is None:
        properties.append(OxmlElement("w:suppressAutoHyphens"))


def _has_math(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))


def _normal(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").strip()).casefold()


def _improve_table_geometry(doc, language: str) -> None:
    rtl = language == "ar"
    section_markers = {
        "lesson plan", "خطة الدرس", "lesson structure", "هيكل الدرس",
    }
    metadata_labels = {
        "teacher:", "subject:", "date:", "class:", "topic:", "periods:",
        "المعلم:", "المادة:", "التاريخ:", "الصف:", "موضوع الدرس:", "عدد الحصص:",
    }

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            unique_cells = []
            seen = set()
            for cell in row.cells:
                if id(cell._tc) not in seen:
                    seen.add(id(cell._tc))
                    unique_cells.append(cell)

            total_text = sum(len(cell.text) for cell in unique_cells)
            is_section_row = any(_normal(cell.text) in section_markers for cell in unique_cells)
            _set_row_keep(row, is_section_row or total_text < 720)

            for cell in unique_cells:
                key = _normal(cell.text)
                if key in metadata_labels:
                    _set_cell_width(cell, 0.92)
                elif key and (key.endswith(":") or key in section_markers) and len(key) < 85:
                    _set_cell_width(cell, 1.35)

                for paragraph in cell.paragraphs:
                    heading = key in section_markers or (
                        paragraph.text.strip().endswith(":") and len(paragraph.text.strip()) < 65
                    )
                    _paragraph_controls(paragraph, heading)
                    if _has_math(paragraph):
                        paragraph.paragraph_format.space_before = Pt(1)
                        paragraph.paragraph_format.space_after = Pt(1)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
                    if rtl:
                        _set_bidi(paragraph, True)


def _remove_orphan_empty_paragraphs(doc) -> None:
    body = doc._element.body
    children = list(body)
    for element in children:
        if element.tag != qn("w:p"):
            continue
        text = "".join(element.itertext()).strip()
        has_drawing = bool(element.xpath(".//w:drawing | .//w:pict"))
        has_break = bool(element.xpath(".//w:br[@w:type='page']"))
        if not text and not has_drawing and not has_break:
            previous = element.getprevious()
            following = element.getnext()
            if previous is not None and following is not None and previous.tag == qn("w:tbl") and following.tag == qn("w:tbl"):
                body.remove(element)


def _document_defaults(doc, language: str) -> None:
    rtl = language == "ar"
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial" if rtl else "Aptos"
        normal.font.size = Pt(8.8)
        normal.paragraph_format.space_after = Pt(1)
        normal.paragraph_format.line_spacing = 1.05
    except Exception:
        pass
    for section in doc.sections:
        section.top_margin = Inches(0.42)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.52)
        section.right_margin = Inches(0.52)
        section.header_distance = Inches(0.06)
        section.footer_distance = Inches(0.12)


def install(word_engine) -> None:
    original_enhance = word_engine.enhance_docx

    def enhanced(data, logger, language="en"):
        first_pass = original_enhance(data, logger, language)
        doc = Document(io.BytesIO(first_pass))
        _document_defaults(doc, language)
        _improve_table_geometry(doc, language)
        _remove_orphan_empty_paragraphs(doc)
        _professional_footer(doc, language)
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    word_engine.enhance_docx = enhanced
